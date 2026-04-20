import re
import fitz                      # PyMuPDF - reads PDF files
from docx import Document        # python-docx - reads DOCX files
import spacy                     # NLP - used to detect candidate name
from langchain_text_splitters import RecursiveCharacterTextSplitter
import config                    # our central settings file



try:
    nlp = spacy.load("en_core_web_sm")
    print("[document_service] spaCy loaded OK")
except OSError:
    print("[document_service] spaCy model missing - run: python -m spacy download en_core_web_sm")
    nlp = None


# ══════════════════════════════════════════════════════════
# STEP 1 - EXTRACT RAW TEXT FROM FILE
# ══════════════════════════════════════════════════════════

def _extract_from_pdf(file_path):
    """Reads a PDF and returns all text as a single string."""
    try:
        text = ""
        pdf = fitz.open(file_path)
        for page in pdf:
            text += page.get_text()
        pdf.close()
        return text
    except Exception as e:
        print(f"[document_service] PDF read error: {e}")
        return None


def _extract_from_docx(file_path):
    """Reads a DOCX file and returns all paragraph and table text as a single string."""
    try:
        doc = Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        # Also extract text from tables (common in resume skill/experience layouts)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        text += cell_text + "\n"
        return text
    except Exception as e:
        print(f"[document_service] DOCX read error: {e}")
        return None


def extract_text(file_path, file_type):
    """Routes extraction to the correct function based on file type."""
    if file_type == "pdf":
        return _extract_from_pdf(file_path)
    elif file_type == "docx":
        return _extract_from_docx(file_path)
    else:
        return None


# ══════════════════════════════════════════════════════════
# STEP 2 - CLEAN THE EXTRACTED TEXT
# ══════════════════════════════════════════════════════════

def clean_text(text):
    """
    Removes noise from extracted text:
    - Fixes common encoding issues (garbled characters)
    - Removes standalone page numbers
    - Collapses excessive whitespace
    """
    if not text:
        return ""

    # Fix common encoding errors that appear when reading PDFs
    text = text.replace("â€¢", "•")
    text = text.replace("â€™", "'")
    text = text.replace("â€œ", '"')
    text = text.replace("â€", '"')
    text = text.replace("\u2019", "'")
    text = text.replace("\u2022", "•")

    # Remove lines that are just page numbers (e.g. "1", "- 2 -", "Page 3")
    lines = text.split("\n")
    cleaned = []
    for line in lines:
        stripped = line.strip()
        if re.match(r'^[-\s]*\d+[-\s]*$', stripped):
            continue
        if re.match(r'^page\s+\d+$', stripped, re.IGNORECASE):
            continue
        cleaned.append(stripped)

    text = "\n".join(cleaned)

    # Collapse multiple spaces and blank lines
    text = re.sub(r' {3,}', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)

    return text.strip()


# ══════════════════════════════════════════════════════════
# STEP 3 - EXTRACT METADATA (name, email, phone)
# ══════════════════════════════════════════════════════════

def extract_metadata(text):
    """
    Pulls out the candidate's basic contact information.
    Uses regex for email and phone (very reliable).
    Uses spaCy NLP for name detection (good but not perfect).
    Falls back to the first clean line if spaCy fails.
    """
    metadata = {
        "full_name": None,
        "email":     None,
        "phone":     None
    }

    # Email - regex is perfect for this pattern
    email_match = re.search(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', text)
    if email_match:
        metadata["email"] = email_match.group()

    # Phone - covers most formats (+94 77 123 4567, 077-123-4567, etc.)
    phone_match = re.search(
        r'(\+?\d{1,3}[\s.-]?)?(\(?\d{2,4}\)?[\s.-]?)(\d{3,4}[\s.-]?\d{3,4})',
        text
    )
    if phone_match:
        metadata["phone"] = phone_match.group().strip()

    # Name - spaCy looks at the first 500 characters (name is always near the top)
    if nlp is not None:
        doc = nlp(text[:500])
        for entity in doc.ents:
            if entity.label_ == "PERSON":
                metadata["full_name"] = entity.text.strip()
                break

    # Fallback: use the first short, clean line that looks like a name
    if not metadata["full_name"]:
        for line in text.split("\n"):
            line = line.strip()
            if not line:
                continue
            if "@" in line:
                continue
            if re.search(r'\d{5,}', line):
                continue
            if line.lower() in ["resume", "curriculum vitae", "cv", "profile"]:
                continue
            if 2 < len(line) < 60:
                metadata["full_name"] = line
                break

    return metadata


# ══════════════════════════════════════════════════════════
# STEP 4 - CHUNK THE TEXT
# ══════════════════════════════════════════════════════════

def chunk_text(text):
    """
    Splits the resume text into smaller pieces (chunks).
    LangChain's splitter tries to break at natural points
    (double newlines, single newlines, sentences) before
    cutting mid-word.

    Chunk size and overlap come from config.py.
    """
    if not text:
        return []

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=config.CHUNK_SIZE,
        chunk_overlap=config.CHUNK_OVERLAP,
        separators=["\n\n", "\n", ". ", " ", ""]
    )
    return splitter.split_text(text)


# ══════════════════════════════════════════════════════════
# STEP 5 - BUILD DOCUMENT STRUCTURE FOR MEMBER 01
# ══════════════════════════════════════════════════════════

def build_documents(candidate_id, job_id, metadata, chunks):
    """
    Wraps each chunk with metadata so Member 01 can store
    it in ChromaDB with full context.

    When Member 01 searches "who has Python experience?",
    ChromaDB returns the matching chunk AND knows which
    candidate and job it belongs to.

    Output shape per chunk:
    {
        "text": "chunk text here...",
        "metadata": {
            "candidate_id": 5,
            "job_id": 2,
            "full_name": "John Smith",
            "email": "john@email.com",
            "phone": "077...",
            "chunk_index": 0
        }
    }
    """
    documents = []
    for index, chunk in enumerate(chunks):
        documents.append({
            "text": chunk,
            "metadata": {
                "candidate_id": candidate_id,
                "job_id":       job_id,
                "full_name":    metadata.get("full_name", "Unknown"),
                "email":        metadata.get("email", ""),
                "phone":        metadata.get("phone", ""),
                "chunk_index":  index
            }
        })
    return documents


# ══════════════════════════════════════════════════════════
# MAIN PIPELINE - runs all steps for one resume file
# ══════════════════════════════════════════════════════════

def process_resume(file_path, file_type, candidate_id, job_id):
    """
    Runs the full processing pipeline for a single resume.
    Called by upload_routes.py.

    Returns a dictionary with:
    - success: True or False
    - metadata: name, email, phone
    - chunks: list of text chunks
    - documents: chunks wrapped with metadata (for Member 01)
    - error: error message if something went wrong
    """

    # Step 1: Extract raw text
    raw_text = extract_text(file_path, file_type)
    if raw_text is None:
        return {"success": False, "error": "Could not read file. It may be corrupted or password protected."}

    # Step 2: Clean the text
    cleaned_text = clean_text(raw_text)
    if not cleaned_text:
        return {"success": False, "error": "File appears to be empty after text extraction."}

    # Step 3: Extract candidate metadata
    metadata = extract_metadata(cleaned_text)

    # Step 4: Split into chunks
    chunks = chunk_text(cleaned_text)

    # Step 5: Wrap chunks with metadata for Member 01
    documents = build_documents(candidate_id, job_id, metadata, chunks)

    return {
        "success":   True,
        "metadata":  metadata,
        "chunks":    chunks,
        "documents": documents,  # this is what Member 01 receives
        "error":     None
    }